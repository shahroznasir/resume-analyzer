import os
import pdfplumber
from docx import Document

def extract_text_from_pdf(file_path: str) -> str:
    pages = []
    try:
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    pages.append(page_text)
        return "\n".join(pages)
    except Exception as e:
        raise Exception(
            f"Failed to extract text from PDF: {str(e)}"
        )

def extract_text_from_docx(file_path: str) -> str:
    try:
        doc = Document(file_path)
        text_parts = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text.strip())
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if paragraph.text.strip():
                            text_parts.append(paragraph.text.strip())
        return "\n".join(text_parts)
    except Exception as e:
        raise Exception(
            f"Failed to extract text from Word document: {str(e)}"
        )

def extract_text_from_txt(file_path: str) -> str:
    try:
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()
    except Exception as e:
        raise Exception(
            f"Failed to extract text from text file: {str(e)}"
        )

def extract_document_text(file_path: str) -> str:
    ext = os.path.splitext(file_path)[1].lower()
    if ext == ".pdf":
        return extract_text_from_pdf(file_path)
    elif ext == ".docx":
        return extract_text_from_docx(file_path)
    elif ext == ".txt":
        return extract_text_from_txt(file_path)
    else:
        raise ValueError(f"Unsupported file extension: {ext}")
